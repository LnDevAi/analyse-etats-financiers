"""
Génération automatisée des rapports d'audit via IA (Claude) + export Word/Excel.
"""
import anthropic
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from typing import Dict, Any, Optional
import os
import uuid
from datetime import datetime

from app.core.config import settings


RISK_COLORS = {
    "VERT": RGBColor(0x16, 0xa3, 0x4a),
    "ORANGE": RGBColor(0xd9, 0x77, 0x06),
    "ROUGE": RGBColor(0xdc, 0x26, 0x26),
}

RISK_HEX = {
    "VERT": "16a34a",
    "ORANGE": "d97706",
    "ROUGE": "dc2626",
}


async def generate_ai_synthesis(
    risk_score: float,
    risk_level: str,
    entity_name: str,
    fiscal_year: Optional[int],
    benford: Optional[Dict],
    isolation_forest: Optional[Dict],
    analytical_review: Optional[Dict],
    cycle_ventes: Optional[Dict],
    cycle_tresorerie: Optional[Dict],
) -> str:
    if not settings.ANTHROPIC_API_KEY:
        return _fallback_synthesis(risk_score, risk_level, entity_name, fiscal_year)

    client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)

    context_parts = [
        f"Entité analysée : {entity_name or 'Non précisé'}",
        f"Exercice fiscal : {fiscal_year or 'Non précisé'}",
        f"Score de confiance global : {risk_score}/100 ({risk_level})",
    ]
    if benford:
        context_parts.append(f"Loi de Benford : conformité {benford.get('conformity_score', 'N/A')}% — {benford.get('risk_level', 'N/A')}")
    if isolation_forest:
        context_parts.append(f"Isolation Forest : {isolation_forest.get('anomalies_detected', 0)} anomalies ({isolation_forest.get('anomaly_rate_pct', 0)}%) — {isolation_forest.get('risk_level', 'N/A')}")
    if cycle_tresorerie:
        context_parts.append(f"Trésorerie : {cycle_tresorerie.get('suspicious_transactions_count', 0)} transactions suspectes — {cycle_tresorerie.get('risk_level', 'N/A')}")
    if cycle_ventes:
        context_parts.append(f"Cut-off ventes : {cycle_ventes.get('cutoff_anomalies_count', 0)} anomalies — {cycle_ventes.get('risk_level', 'N/A')}")

    prompt = f"""Tu es un Expert-Comptable et Commissaire aux Comptes senior, spécialiste SYSCOHADA.
Rédige une note de synthèse d'audit professionnelle et concise (300-400 mots maximum) en français.

Données de l'analyse automatisée :
{chr(10).join(context_parts)}

La note doit :
1. Rappeler l'objet et le périmètre de l'analyse
2. Présenter les principales observations par module (Benford, anomalies ML, trésorerie, cut-off)
3. Formuler une conclusion sur le niveau de risque avec recommandations concrètes
4. Adopter un ton professionnel et factuel, conforme aux normes ISA/SYSCOHADA

NE PAS inventer de chiffres non fournis. Se baser uniquement sur les données ci-dessus."""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text


def _fallback_synthesis(score: float, level: str, entity: str, year: Optional[int]) -> str:
    return (
        f"NOTE DE SYNTHÈSE D'AUDIT AUTOMATISÉE\n\n"
        f"Entité : {entity or 'Non précisé'} | Exercice : {year or 'N/A'}\n"
        f"Score global de confiance : {score}/100 — Niveau : {level}\n\n"
        f"L'analyse automatisée des états financiers par les algorithmes E-DÉFENCE SaaS "
        f"a produit un score de confiance de {score}/100, classant cet exercice en niveau {level}.\n\n"
        f"Voir le détail par module dans le rapport complet."
    )


def generate_word_report(
    analysis_data: Dict[str, Any],
    entity_name: str,
    fiscal_year: Optional[int],
    output_dir: str,
) -> str:
    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # En-tête
    header = doc.add_heading("E-DÉFENCE — RAPPORT D'ANALYSE FINANCIÈRE IA", level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    doc.add_paragraph(f"Entité : {entity_name or 'Non précisé'}")
    doc.add_paragraph(f"Exercice fiscal : {fiscal_year or 'Non précisé'}")
    doc.add_paragraph(f"Date du rapport : {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("Confidentiel — Usage interne uniquement").runs[0].italic = True

    doc.add_heading("1. SCORE DE RISQUE GLOBAL", level=2)
    score = analysis_data.get("risk_score", 0)
    level = analysis_data.get("risk_level", "N/A")
    p = doc.add_paragraph(f"Score de confiance : {score}/100 — Niveau de risque : {level}")
    if level in RISK_COLORS:
        p.runs[0].font.color.rgb = RISK_COLORS[level]
        p.runs[0].bold = True

    if analysis_data.get("ai_synthesis"):
        doc.add_heading("2. NOTE DE SYNTHÈSE", level=2)
        doc.add_paragraph(analysis_data["ai_synthesis"])

    modules = [
        ("3. LOI DE BENFORD", "benford_result"),
        ("4. ANOMALIES ML (ISOLATION FOREST)", "isolation_forest_result"),
        ("5. REVUE ANALYTIQUE N vs N-1", "analytical_review"),
        ("6. CYCLE VENTES / CUT-OFF", "cycle_ventes_result"),
        ("7. CYCLE TRÉSORERIE", "cycle_tresorerie_result"),
    ]

    for title, key in modules:
        data = analysis_data.get(key)
        if not data:
            continue
        doc.add_heading(title, level=2)
        rl = data.get("risk_level", "N/A")
        p = doc.add_paragraph(f"Niveau : {rl}")
        if rl in RISK_COLORS:
            p.runs[0].font.color.rgb = RISK_COLORS[rl]
        interp = data.get("interpretation") or data.get("message", "")
        if interp:
            doc.add_paragraph(interp)

    doc.add_heading("MENTIONS LÉGALES", level=3)
    doc.add_paragraph(
        "Ce rapport a été généré automatiquement par E-DÉFENCE SaaS V4. "
        "Il constitue une aide à la décision et ne remplace pas le jugement professionnel "
        "d'un expert-comptable ou commissaire aux comptes certifié."
    ).italic = True

    filename = f"rapport_audit_{uuid.uuid4().hex[:8]}.docx"
    path = os.path.join(output_dir, filename)
    os.makedirs(output_dir, exist_ok=True)
    doc.save(path)
    return path


def generate_excel_report(
    analysis_data: Dict[str, Any],
    entity_name: str,
    fiscal_year: Optional[int],
    output_dir: str,
) -> str:
    wb = openpyxl.Workbook()
    ws_summary = wb.active
    ws_summary.title = "Synthèse"

    header_fill = PatternFill("solid", fgColor="1e293b")
    header_font = Font(color="FFFFFF", bold=True)
    bold_font = Font(bold=True)

    ws_summary["A1"] = "E-DÉFENCE — ANALYSE FINANCIÈRE IA"
    ws_summary["A1"].font = Font(bold=True, size=14, color="1e293b")
    ws_summary["A2"] = f"Entité : {entity_name or 'N/A'}"
    ws_summary["A3"] = f"Exercice : {fiscal_year or 'N/A'}"
    ws_summary["A4"] = f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}"

    ws_summary["A6"] = "Module"
    ws_summary["B6"] = "Niveau de risque"
    ws_summary["C6"] = "Score"
    ws_summary["D6"] = "Interprétation"

    for cell in [ws_summary["A6"], ws_summary["B6"], ws_summary["C6"], ws_summary["D6"]]:
        cell.fill = header_fill
        cell.font = header_font

    row = 7
    modules_data = [
        ("Score Global", analysis_data.get("risk_level"), analysis_data.get("risk_score"), ""),
        ("Vérification Intrinsèque", (analysis_data.get("intrinsic_check") or {}).get("risk_level"), None, ""),
        ("Loi de Benford", (analysis_data.get("benford_result") or {}).get("risk_level"), (analysis_data.get("benford_result") or {}).get("conformity_score"), (analysis_data.get("benford_result") or {}).get("interpretation", "")),
        ("Isolation Forest", (analysis_data.get("isolation_forest_result") or {}).get("risk_level"), None, (analysis_data.get("isolation_forest_result") or {}).get("interpretation", "")),
        ("Revue N vs N-1", (analysis_data.get("analytical_review") or {}).get("risk_level"), None, ""),
        ("Cycle Ventes", (analysis_data.get("cycle_ventes_result") or {}).get("risk_level"), None, (analysis_data.get("cycle_ventes_result") or {}).get("interpretation", "")),
        ("Cycle Trésorerie", (analysis_data.get("cycle_tresorerie_result") or {}).get("risk_level"), None, ""),
    ]

    for name, level, score, interp in modules_data:
        ws_summary[f"A{row}"] = name
        ws_summary[f"B{row}"] = level or "N/A"
        ws_summary[f"C{row}"] = score or ""
        ws_summary[f"D{row}"] = str(interp)[:200] if interp else ""

        if level in RISK_HEX:
            fill = PatternFill("solid", fgColor=RISK_HEX[level])
            ws_summary[f"B{row}"].fill = fill
            ws_summary[f"B{row}"].font = Font(color="FFFFFF", bold=True)

        row += 1

    for col in ["A", "B", "C", "D"]:
        ws_summary.column_dimensions[col].width = 30 if col == "D" else 20

    filename = f"rapport_audit_{uuid.uuid4().hex[:8]}.xlsx"
    path = os.path.join(output_dir, filename)
    os.makedirs(output_dir, exist_ok=True)
    wb.save(path)
    return path
