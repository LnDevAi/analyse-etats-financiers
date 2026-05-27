"""
Extraction de texte et de tableaux depuis différents formats de documents AG.
Supporte : PDF (via pdfplumber), Excel (openpyxl), CSV, Word (.docx).
"""
import io
import re
import logging
import pandas as pd
from typing import Dict, Any, List, Optional, Tuple
import chardet

logger = logging.getLogger(__name__)

# Montants FCFA / XOF dans du texte : "1 234 567 FCFA", "2,5 milliards", "450 millions"
AMOUNT_PATTERN = re.compile(
    r"""
    (?:
        (\d[\d\s,\.]*\d|\d)          # Nombre (avec séparateurs)
        \s*
        (?:
            (milliard|million|millier)s?  # Unité textuelle
            (?:\s+(?:de\s+)?)?
            (?:FCFA|XOF|F\.?CFA|CFA)?
        |
            (?:FCFA|XOF|F\.?CFA|CFA)  # Unité monétaire directe
        )
    )
    """,
    re.VERBOSE | re.IGNORECASE,
)

MULTIPLIERS = {"milliard": 1_000_000_000, "million": 1_000_000, "millier": 1_000}


def _clean_number(s: str) -> float:
    """Nettoie une chaîne numérique : supprime espaces, convertit virgule en point."""
    s = s.strip().replace("\xa0", "").replace(" ", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return 0.0


def extract_amounts_from_text(text: str) -> List[Dict[str, Any]]:
    """
    Extrait tous les montants financiers présents dans un texte.
    Retourne une liste {raw, value, context}.
    """
    results = []
    for match in AMOUNT_PATTERN.finditer(text):
        raw = match.group(0).strip()
        num_str = match.group(1) or "0"
        unit = (match.group(2) or "").lower()
        value = _clean_number(num_str) * MULTIPLIERS.get(unit, 1)
        if value > 0:
            start = max(0, match.start() - 60)
            end = min(len(text), match.end() + 60)
            context = text[start:end].replace("\n", " ").strip()
            results.append({"raw": raw, "value": value, "context": context})
    return results


def _detect_encoding(content: bytes) -> str:
    r = chardet.detect(content[:10000])
    return r.get("encoding") or "utf-8"


def extract_from_pdf(content: bytes) -> Tuple[str, List[pd.DataFrame]]:
    """Extrait texte + tableaux d'un PDF via pdfplumber."""
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        logger.warning("pdfplumber non installé — extraction PDF limitée au texte brut.")
        return _extract_pdf_fallback(content)

    full_text = ""
    tables: List[pd.DataFrame] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                full_text += page_text + "\n"
                for table in page.extract_tables():
                    if table and len(table) > 1:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        df.columns = [str(c).strip() if c else f"col_{i}" for i, c in enumerate(df.columns)]
                        tables.append(df)
    except Exception as e:
        logger.error(f"Erreur extraction PDF : {e}")
    return full_text, tables


def _extract_pdf_fallback(content: bytes) -> Tuple[str, List[pd.DataFrame]]:
    """Fallback minimal si pdfplumber absent — lecture raw bytes."""
    try:
        text = content.decode(_detect_encoding(content), errors="replace")
        # Filtre les caractères non-imprimables
        text = re.sub(r"[^\x20-\x7EÀ-ž\n\r\t]", " ", text)
        return text, []
    except Exception:
        return "", []


def extract_from_excel(content: bytes) -> Tuple[str, List[pd.DataFrame]]:
    """Extrait texte + feuilles d'un fichier Excel."""
    tables = []
    texts = []
    try:
        xl = pd.ExcelFile(io.BytesIO(content))
        for sheet in xl.sheet_names:
            df = xl.parse(sheet, dtype=str)
            tables.append(df)
            texts.append(f"[Feuille: {sheet}]\n{df.to_string(index=False)}")
    except Exception as e:
        logger.error(f"Erreur extraction Excel : {e}")
    return "\n".join(texts), tables


def extract_from_csv(content: bytes) -> Tuple[str, List[pd.DataFrame]]:
    """Extrait un CSV en DataFrame + texte."""
    enc = _detect_encoding(content)
    text = content.decode(enc, errors="replace")
    first_line = text.split("\n")[0]
    sep = "\t" if "\t" in first_line else (";" if ";" in first_line else ",")
    try:
        df = pd.read_csv(io.StringIO(text), sep=sep, dtype=str, on_bad_lines="skip")
        return df.to_string(index=False), [df]
    except Exception as e:
        logger.error(f"Erreur CSV : {e}")
        return text, []


def extract_from_docx(content: bytes) -> Tuple[str, List[pd.DataFrame]]:
    """Extrait texte + tableaux d'un fichier Word."""
    try:
        from docx import Document  # type: ignore
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for tbl in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            if rows:
                df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
                tables.append(df)
        return "\n".join(paragraphs), tables
    except Exception as e:
        logger.error(f"Erreur Word : {e}")
        return "", []


def extract_document(content: bytes, filename: str) -> Dict[str, Any]:
    """
    Route l'extraction selon l'extension du fichier.
    Retourne : {text, tables, amounts, filename}
    """
    fname = filename.lower()
    if fname.endswith(".pdf"):
        text, tables = extract_from_pdf(content)
    elif fname.endswith((".xlsx", ".xls")):
        text, tables = extract_from_excel(content)
    elif fname.endswith((".docx",)):
        text, tables = extract_from_docx(content)
    elif fname.endswith((".csv", ".txt", ".tsv")):
        text, tables = extract_from_csv(content)
    else:
        # Tentative CSV par défaut
        text, tables = extract_from_csv(content)

    amounts = extract_amounts_from_text(text)
    return {
        "filename": filename,
        "text_length": len(text),
        "text_preview": text[:500],
        "tables_found": len(tables),
        "tables": tables,
        "amounts_found": amounts,
    }
